import speech_recognition as sr
import pyttsx3
import pyautogui
import time
import wikipedia
import os
import docx
#from docx import Document

engine = pyttsx3.init()
engine.setProperty('rate', 150)

def speak(audio):
    #spee-out what is present in audio string
    engine.say(audio)
    engine.runAndWait()

def open_document(doc_name):
    if doc_name:
        if os.path.exists(f"{doc_name}.docx"):
            os.startfile(f"{doc_name}.docx")
        else:
            speak("Document not found.")
            print("Document not found.")

def close_document():
    try:
        speak("saving closing the document")
        time.sleep(1)
        pyautogui.hotkey('ctrl', 's')
        pyautogui.hotkey('ctrl', 'w')
        print("Microsoft Word is closed.")
    except Exception as e:
        print(f"Error occurred while closing Word: {e}")

def create_document(doc_name):
    if doc_name:
        doc = Document()
        doc.save(f"{doc_name}.docx")
        print(f"Document '{doc_name}.docx' created.")
        speak("document created")

def speech_to_text():
    recognizer = sr.Recognizer()

    with sr.Microphone() as source:
        print("Speak something...")
        recognizer.adjust_for_ambient_noise(source)
        audio = recognizer.listen(source)
    try:
        print("Transcribing...")
        text= recognizer.recognize_google(audio, language="en-US")
        return text
    except sr.UnknownValueError:
        print("Google Speech Recognition could not understand the audio.")
    except sr.RequestError as e:
        print(f"Could not request results from Google Speech Recognition service; {e}")

def write_to_word(text):
    pyautogui.write(text)

def start_writing():
    while True:
        #print("speech...")
        speech_text = speech_to_text()
        if speech_text=="stop writing":
            speak("writing stopped")
            break
        else:
            speech_text=speech_text+" "
            write_to_word(speech_text)
            pyautogui.hotkey('ctrl', 's')

def search_wiki():
    query=speech_to_text()
    result = wikipedia.summary(query,sentences=4)
    write_to_word(result)

def read(file_path):
    try:
        doc = docx.Document(file_path)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        speak(text)
    except Exception as e:
        print("Error:", e)
        speak("unable to read data")
        return None


if __name__ == "__main__":
    speak("welcome i am your document bot assistent you can call me jvs")
    time.sleep(1)
    speak("i help you in convertion of speech to text and write it in a document")
    time.sleep(1)
    speak("i can accept english language")
    time.sleep(1)
    speak("opening ms word")
    time.sleep(1)
    speak("say command start writing to start the writing process")
    document_name=None
    while True:
        print("listenning cmd...")
        cmd=speech_to_text()
        if "start writing" in cmd:
            if not document_name:
                speak("Please create or open a document first.")
            else:
                start_writing()
    
        elif "open document" in cmd:
            speak("What is the name of the document you want to open?")
            document_name = speech_to_text()
            open_document(document_name)
        
        elif "close document" in cmd:
            close_document()
        
        elif "create document" in cmd:
            speak("What do you want to name the document?")
            document_name = speech_to_text()
            create_document(document_name)
        
        elif "read" in cmd:
            speak("What is the name of the document you want to read?")
            document_name = speech_to_text()
            if document_name:
                # Open the file in read mode ('r')
                open_document(document_name)
                document_name=document_name+".docx"
                print(document_name)
                read(document_name)
            else:
                print("Document does not exist")


        elif "search" in cmd:
            search_wiki()
